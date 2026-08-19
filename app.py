#!/usr/bin/env python3
"""
app.py — Local web app: paste a URL, see its sublinks, download all as PDFs.
Run:  python app.py
Open: http://localhost:5000
"""

import asyncio, hashlib, html, json, os, queue, re, tempfile, threading, time, uuid, zipfile
from datetime import datetime
from pathlib import Path
from urllib.parse import urljoin, urlparse
from urllib.request import Request, urlopen

import anthropic
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.opc.constants
import openpyxl
from dotenv import load_dotenv
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from flask import Flask, Response, jsonify, render_template, request, send_file
from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeout
from playwright.async_api import async_playwright
from classifier import (
    HEALTH_DIMENSIONS,
    clamp_analysis_fields,
    default_analysis_fields,
)

load_dotenv()
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# ── Content Management Guide → cached Claude system prompt ─────────────────────
# The guide markdown is the single source of truth.  We extract its
# platform-definitions section and inject it verbatim into the system prompt so
# that editing the guide changes the AI's behavior with no code changes.
GUIDE_PATH = Path("University_Comprehensive_Content_Management_Guide.md")
GUIDE_SECTION_START = "## Overview of MSU's Main Content Platforms"
GUIDE_SECTION_END = "## Content Audit"


def load_platform_guide() -> str:
    """Extract the platform-definitions section from the content management guide."""
    try:
        text = GUIDE_PATH.read_text(encoding="utf-8")
    except Exception:
        return ""
    start = text.find(GUIDE_SECTION_START)
    if start == -1:
        return ""
    end = text.find(GUIDE_SECTION_END, start + len(GUIDE_SECTION_START))
    return (text[start:end] if end != -1 else text[start:]).strip()


def guide_version() -> str:
    """Short hash of the guide file, used to invalidate cached suggestions on edit."""
    try:
        return hashlib.md5(GUIDE_PATH.read_bytes()).hexdigest()[:8]
    except Exception:
        return "noguide"


# ── AI Content Analysis framework → cached Claude system prompt ────────────────
# Injected whole rather than sliced: its batch-level sections (17-18) are a small
# fraction of the file, and TIER2_INSTRUCTIONS tells the model to skip them.
FRAMEWORK_PATH = Path("AI_Content_Analysis_Direction_Support.md")


def load_ai_analysis_framework() -> str:
    """Load the AI content-analysis framework verbatim."""
    try:
        return FRAMEWORK_PATH.read_text(encoding="utf-8").strip()
    except Exception:
        return ""


def framework_version() -> str:
    """Short hash of the framework file, used to invalidate cached suggestions on edit."""
    try:
        return hashlib.md5(FRAMEWORK_PATH.read_bytes()).hexdigest()[:8]
    except Exception:
        return "noframework"


# Fixed instructions; the guide and framework blocks are appended and prompt-cached.
TIER2_INSTRUCTIONS = """You are a content strategist and content-UX analyst at Minnesota State University, Mankato (MNSU).

Your job: read ONE university web page — given its title, URL, and a content excerpt — and produce a complete, structured analysis of that single page. Two source documents follow below: the official Content Management Guide (platform definitions) and the AI Content Analysis & Platform Recommendation Framework (how to evaluate a page). Base every judgment strictly on those two documents, not on general web-content-strategy knowledge.

You are evaluating pages ONE AT A TIME. The framework document's sections 17 (Cross-Site Analysis) and 18 (Project Plan Generation) describe a SEPARATE, batch-level step that runs after every page in a crawl has been analyzed — they are not part of this task. Ignore them; do not attempt cross-page comparisons beyond what the current excerpt tells you.

Judge by the actual content and its primary audience and purpose — NOT by the URL structure. A page sitting under a public-facing section can still belong on another platform.

═══ PLATFORM RECOMMENDATION ═══

VALID PLATFORMS (respond with one of these EXACT names):
- Website
- Maverick OneStop
- The Fountain
- MavLife / Student Hub
- Teams / SharePoint
- No Clear Fit

DECISION GUIDANCE (always defer to the Content Management Guide below):
- Policies, conduct codes, required procedures, forms, FAQs, and step-by-step task instructions → Maverick OneStop, even under a public URL like /housing/policies/.
- Purely informational or marketing pages for prospective students, families, or the general public → Website.
- Content exclusively for employees (not students) → The Fountain.
- Student involvement, clubs, activities, recreation, engagement → MavLife / Student Hub.
- Internal department/committee collaboration material → Teams / SharePoint.
- If nothing genuinely fits, use "No Clear Fit".

Every page you see currently lives on the public Website — that is where it was crawled from. Only name a DIFFERENT platform when the content genuinely belongs elsewhere.

═══ FULL PAGE ANALYSIS ═══

1. USER NEED — primary: Explore | Understand | Act | Connect. Secondary: the same four, or "Not applicable".
2. AUDIENCE — primary: Prospective students | Current students | Families/supporters | Faculty | Staff | Alumni | Community/public | Multiple audiences | Unable to determine. Secondary: the same nine, or "Not applicable".
3. CONTENT PURPOSE — one of: Recruitment | Awareness | General information | Service information | Process guidance | Transaction support | Policy or requirement | Academic information | Event or engagement | News or announcement | Contact/support | Resource collection | Reference information | Other. Add one short sentence explaining it.
4. CONTENT HEALTH — rate EACH of six dimensions independently as Strong | Acceptable | Needs Improvement | Significant Concern | Unable to Determine:
   - Accuracy and Currency (outdated dates, expired deadlines, stale terminology — do NOT flag evergreen content merely for lacking a date)
   - Clarity (dense text, jargon, unclear instructions)
   - Scanability and Structure (headings, lists, logical sectioning)
   - Actionability (clear next steps, calls to action, contact options)
   - User Focus (organized around user needs vs. internal department structure)
   - Sustainability (manually-repeated info that will rot vs. content easy to keep current)
   For every dimension rated Needs Improvement or Significant Concern, add one short clause to "health_notes" naming the dimension and the specific issue (e.g. "Clarity: the reporting steps are one dense paragraph."). Combine all clauses into a single string; use "" if none apply.
5. DUPLICATION — one of: No significant duplication found | Possible overlap | Significant overlap | Likely duplicate | Conflicting content detected. Put up to 5 related URLs in "related_urls" ONLY with genuine evidence from the excerpt; otherwise [].
6. AI/SEARCH READINESS — how well this page would serve as a source for AI-assisted answers (e.g. MNSU's "Ask Stomper"). Rate on the same five-value scale as content health.
7. RECOMMENDED TREATMENT — exactly one of: KEEP | KEEP + IMPROVE | CONSOLIDATE | CONNECT | CONSIDER MOVING | REPLACE WITH ACTION | ARCHIVE / REMOVE REVIEW | EXPERT REVIEW NEEDED.
8. IMPROVEMENT OPPORTUNITIES — up to 5 from: Update outdated information, Verify accuracy, Clarify primary audience, Clarify page purpose, Rewrite for clarity, Reduce jargon, Improve headings, Improve page title, Improve scanability, Add clear next step, Improve call to action, Clarify eligibility, Clarify deadlines, Improve contact/support information, Consolidate duplicate content, Establish authoritative source, Remove redundant information, Link rather than duplicate, Improve platform fit, Improve AI/search readability, Separate multiple audiences, Improve user journey, Review for removal — or just "No significant improvement identified".
9. PRIORITY — one of: "Priority 1 - Act Now" | "Priority 2 - Important Improvement" | "Priority 3 - Strategic Opportunity" | "Priority 4 - Maintain/Monitor".
10. SUMMARY FIELDS — "page_summary" (1-2 sentences on what the page provides), "major_issues" (1-2 sentences on the most important strength or concern), "key_recommendation" (one sentence on what should happen), "content_owner_questions" (up to 3 questions genuinely needing a human answer, else []).

CONFIDENCE CALIBRATION — be honest, never manufacture confidence:
- High: The content unambiguously fits ONE platform and treatment; an expert would agree instantly.
- Medium: One answer is best, but a defensible alternative exists, or institutional context is missing.
- Low: Multiple answers are genuinely defensible, or the excerpt is too thin. Use this freely.
When confidence is not High, state in "confidence_gap" what additional information would raise it; use "" when it is High.

GUARDRAILS — you must NOT:
- Assume every page needs to change. KEEP is a valid, common outcome.
- Treat short content as automatically good, or long content as automatically bad.
- Invent traffic, search volume, user behavior, conversion rates, ownership, institutional requirements, analytics, or business rules. Say "Unable to Determine" instead.
- Recommend or imply automatic deletion. ARCHIVE / REMOVE REVIEW always means "flag for human review".
- Treat Ask Stomper (or any AI assistant) as a content destination.
- Reinterpret policy language in a way that changes its legal or compliance meaning. Comment on the clarity and structure of policy content, not its substance.

RESPONSE FORMAT — a single raw JSON object, no markdown, no code fences, no text before or after. Always include EVERY key below, using these exact names. Use the literal defaults ("Not applicable", "", []) when a field genuinely does not apply — never omit a key:

{
  "platform": "<one of the 6 valid platforms>",
  "confidence": "High|Medium|Low",
  "reason": "<one concise sentence grounding the platform choice in the guide>",
  "page_summary": "<1-2 sentences>",
  "primary_user_need": "Explore|Understand|Act|Connect",
  "secondary_user_need": "Explore|Understand|Act|Connect|Not applicable",
  "primary_audience": "<one of the 9 audience values>",
  "secondary_audience": "<one of the 9 audience values, or Not applicable>",
  "content_purpose": "<one of the 14 purpose values>",
  "purpose_explanation": "<short sentence>",
  "health_accuracy_currency": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "health_clarity": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "health_scanability_structure": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "health_actionability": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "health_user_focus": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "health_sustainability": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "health_notes": "<combined dimension: issue clauses, or empty string>",
  "duplication_status": "<one of the 5 duplication values>",
  "related_urls": ["<up to 5 URLs>"],
  "ai_search_readiness": "Strong|Acceptable|Needs Improvement|Significant Concern|Unable to Determine",
  "recommended_treatment": "<one of the 8 treatment values>",
  "improvement_opportunities": ["<up to 5 improvement values>"],
  "priority": "Priority 1 - Act Now|Priority 2 - Important Improvement|Priority 3 - Strategic Opportunity|Priority 4 - Maintain/Monitor",
  "confidence_gap": "<what would raise confidence, or empty string>",
  "major_issues": "<1-2 sentences>",
  "key_recommendation": "<one sentence>",
  "content_owner_questions": ["<up to 3 short questions, or empty list>"]
}

--- OFFICIAL CONTENT MANAGEMENT GUIDE (SOURCE OF TRUTH FOR PLATFORM DEFINITIONS) ---
"""

PLATFORM_GUIDE = load_platform_guide()
GUIDE_VERSION = guide_version()
AI_ANALYSIS_FRAMEWORK = load_ai_analysis_framework()
FRAMEWORK_VERSION = framework_version()

FRAMEWORK_BLOCK_TEXT = (
    "--- AI CONTENT ANALYSIS & PLATFORM RECOMMENDATION FRAMEWORK "
    "(SOURCE OF TRUTH FOR THE ANALYSIS DIMENSIONS ABOVE; IGNORE ITS SECTIONS "
    "17-18, THOSE ARE A SEPARATE BATCH-LEVEL STEP) ---\n\n" + AI_ANALYSIS_FRAMEWORK
)

# Anthropic prompt-caching: the marker sits on the LAST block so the whole
# constant prefix (instructions + guide + framework) is cached as one unit
# across the batch, cutting input cost and latency.
TIER2_SYSTEM_BLOCKS = [
    {"type": "text", "text": TIER2_INSTRUCTIONS},
    {"type": "text", "text": PLATFORM_GUIDE},
    {"type": "text", "text": FRAMEWORK_BLOCK_TEXT, "cache_control": {"type": "ephemeral"}},
]

# Bump when TIER2_INSTRUCTIONS or the response schema changes, so cached results
# in the old shape can never be served as if they were current.
SUGGESTION_SCHEMA_VERSION = "2"
SUGGEST_CACHE_VERSION = f"{GUIDE_VERSION}-{FRAMEWORK_VERSION}-{SUGGESTION_SCHEMA_VERSION}"

# Concurrency + model settings for the AI analysis pipeline.
# ANALYSIS_CONCURRENCY caps how many pages are fetched + classified at once.
# Each in-flight page is a live Chromium tab, so this is the main CPU/RAM dial.
# Default 3 keeps a typical laptop responsive; raise on a beefy machine, lower
# (e.g. 2) on memory-constrained hosts like the Render free tier (512 MB RAM).
TIER2_MODEL = "claude-haiku-4-5"
AI_CREDIT_ERROR_MARKERS = (
    "credit balance",
    "insufficient credit",
    "insufficient_quota",
    "billing",
    "payment required",
    "quota exceeded",
)


class AIQuotaError(Exception):
    pass


def env_int(name: str, default: int, minimum: int | None = None) -> int:
    try:
        value = int(os.getenv(name, str(default)))
    except (TypeError, ValueError):
        value = default
    if minimum is not None:
        value = max(value, minimum)
    return value


ANALYSIS_CONCURRENCY = env_int("ANALYSIS_CONCURRENCY", 2, 1)

# Response budget for one page's analysis JSON.  Worst-case realistic output is
# ~800 tokens; the headroom guards against truncated (unparseable) responses.
TIER2_MAX_TOKENS = env_int("TIER2_MAX_TOKENS", 2048, 256)

# Restart the shared Chromium browser every N page loads.  A single long-lived
# context accumulates memory across thousands of navigations, which is what
# freezes machines on large (~2000 link) runs.  Recycling caps peak RAM.
ANALYSIS_BROWSER_RECYCLE = env_int("ANALYSIS_BROWSER_RECYCLE", 100, 1)

# Per-page navigation timeout (ms) for the analysis fetch.
ANALYSIS_NAV_TIMEOUT = env_int("ANALYSIS_NAV_TIMEOUT", 15000, 1000)
CRAWL_NAV_TIMEOUT = env_int("CRAWL_NAV_TIMEOUT", 10000, 1000)
PDF_NAV_TIMEOUT = env_int("PDF_NAV_TIMEOUT", 20000, 1000)
PDF_SETTLE_MS = env_int("PDF_SETTLE_MS", 750, 0)
PDF_PAGE_RECYCLE = env_int("PDF_PAGE_RECYCLE", 75, 1)
SSE_HEARTBEAT_SECONDS = env_int("SSE_HEARTBEAT_SECONDS", 20, 5)
MAX_ACTIVE_JOBS = env_int("MAX_ACTIVE_JOBS", 2, 1)
JOB_RETENTION_SECONDS = env_int("JOB_RETENTION_SECONDS", 3600, 60)

DATA_DIR = Path(os.getenv("DATA_DIR", tempfile.gettempdir() if os.getenv("RAILWAY_ENVIRONMENT") else "."))
SUGGEST_CACHE_DIR = DATA_DIR / "suggestion_cache"

# Safety ceiling on how many pages a single crawl will visit.  The BFS normally
# stops on its own once it exhausts the in-scope queue; this only caps runaway
# crawls.  Default is well above large sections so full coverage isn't truncated.
CRAWL_MAX_PAGES = env_int("CRAWL_MAX_PAGES", 1000, 1)
SAFE_PATH_SEGMENT_MAX = env_int("SAFE_PATH_SEGMENT_MAX", 48, 16)
SAFE_REL_PATH_MAX = env_int("SAFE_REL_PATH_MAX", 150, 80)
print(f"[startup] Python 3", flush=True)
for mod in ["docx", "anthropic", "openpyxl", "flask", "gunicorn", "playwright"]:
    try:
        __import__(mod)
        print(f"[startup] {mod}: ok", flush=True)
    except Exception as e:
        print(f"[startup] {mod}: FAIL - {e}", flush=True)
print("[startup] app.py loaded, binding...", flush=True)


# Site chrome that should never appear in an exported PDF/DOCX.  MNSU wraps its
# yellow footer in .footer-background/.footer-wrap and the purple legal strip in
# footer.sub-footer, so those containers are targeted explicitly.  The same list
# drives the injected CSS (for the PDF) and the DOCX text extractor, so the two
# formats can never disagree about what counts as boilerplate.
BOILERPLATE_SELECTORS = ", ".join([
    "header", "nav", ".nav", ".navbar", ".site-header",
    "footer", ".footer", ".site-footer", ".sub-footer",
    ".footer-background", ".footer-wrap",
    ".foot-nav", ".foot-block", ".foot-text",
    ".cookie-banner", ".chat-widget", "[class*=\"overlay\"]",
    "script", "style", "noscript",
])

PAGE_CLEANUP_CSS = """
""" + BOILERPLATE_SELECTORS + """ {
    display: none !important;
}
img, svg, video, audio, canvas,
[class*="icon"], [class*="logo"], [class*="banner"],
[class*="hero"], [class*="thumbnail"], [class*="carousel"],
[class*="slider"], [class*="gallery"], [class*="image"],
picture, figure, iframe {
    display: none !important;
}
* {
    background-image: none !important;
}
[style*="position: fixed"], [style*="position:fixed"],
[style*="position: sticky"], [style*="position:sticky"] {
    position: static !important;
}
p, h1, h2, h3, h4, h5, h6, table, li, blockquote {
    page-break-inside: avoid;
    break-inside: avoid;
}
"""

app = Flask(__name__)

# ── CORS ─────────────────────────────────────────────────────────────────────
# Allow any origin so a separately-deployed frontend (e.g. on Vercel) can call
# this backend and open SSE streams.  No credentials are used, so * is safe.
ALLOWED_ORIGIN = os.getenv("ALLOWED_ORIGIN", "*")

@app.after_request
def add_cors_headers(response):
    response.headers["Access-Control-Allow-Origin"]  = ALLOWED_ORIGIN
    response.headers["Access-Control-Allow-Headers"] = "Content-Type"
    response.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return response

@app.route("/", defaults={"path": ""}, methods=["OPTIONS"])
@app.route("/<path:path>", methods=["OPTIONS"])
def handle_preflight(path):
    """Return 200 for every CORS preflight so fetch() with JSON bodies works cross-origin."""
    resp = app.make_default_options_response()
    resp.headers["Access-Control-Allow-Origin"]  = ALLOWED_ORIGIN
    resp.headers["Access-Control-Allow-Headers"] = "Content-Type"
    resp.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
    return resp
jobs: dict = {}          # job_id -> {"type": "scrape"|"download", "queue": Queue, "links": list|None, "zip_path": str|None}
jobs_lock = threading.Lock()

OUTPUT_DIR = DATA_DIR / "mankato_pdfs"
MANIFEST_PATH = OUTPUT_DIR / "manifest.json"
CRAWL_CACHE_DIR = DATA_DIR / "crawl_cache"  # Cache crawl results by starting URL


def cleanup_jobs() -> None:
    cutoff = time.time() - JOB_RETENTION_SECONDS
    with jobs_lock:
        for job_id, job in list(jobs.items()):
            if job.get("finished_at", 0) and job["finished_at"] < cutoff:
                jobs.pop(job_id, None)


def register_job(job_id: str, job: dict) -> bool:
    cleanup_jobs()
    with jobs_lock:
        active = sum(1 for item in jobs.values() if not item.get("finished_at"))
        if active >= MAX_ACTIVE_JOBS:
            return False
        job["created_at"] = time.time()
        job["cancelled"] = False
        jobs[job_id] = job
        return True


def should_stop(job_id: str) -> bool:
    with jobs_lock:
        job = jobs.get(job_id)
        return bool(job and job.get("cancelled"))


def cancel_job(job_id: str) -> bool:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job or job.get("finished_at"):
            return False
        job["cancelled"] = True
        job.setdefault("status", {}).update(stage="cancelling", updated_at=time.time())
        q = job.get("queue")
    if q:
        q.put({"type": "cancelling"})
    return True


def timestamp_slug() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def update_job_status(job_id: str, **status) -> None:
    with jobs_lock:
        job = jobs.get(job_id)
        if job:
            current = job.setdefault("status", {})
            current.update(status, updated_at=time.time())


def log_job(job_id: str, message: str, **details) -> None:
    detail_text = " ".join(f"{key}={value}" for key, value in details.items())
    print(f"[job:{job_id}] {message} {detail_text}".rstrip(), flush=True)


def finish_job(job_id: str) -> None:
    with jobs_lock:
        job = jobs.get(job_id)
        if job and not job.get("finished_at"):
            job["finished_at"] = time.time()


def emit_terminal(q: queue.Queue, job_id: str, event: dict) -> None:
    update_job_status(job_id, stage=event.get("type"), terminal=True)
    with jobs_lock:
        job = jobs.get(job_id)
        if job is not None:
            job["terminal"] = event
    finish_job(job_id)
    q.put(event)


HEAVY_RESOURCE_TYPES = {"image", "media", "font"}
BLOCKED_URL_PARTS = ("googletagmanager", "google-analytics", "doubleclick", "hotjar", "fullstory")


def should_block_resource(request, include_stylesheets: bool = False) -> bool:
    if request.resource_type in HEAVY_RESOURCE_TYPES:
        return True
    if include_stylesheets and request.resource_type == "stylesheet":
        return True
    return any(part in request.url.lower() for part in BLOCKED_URL_PARTS)


def block_resource_route(route, include_stylesheets: bool = False) -> None:
    if should_block_resource(route.request, include_stylesheets):
        route.abort()
    else:
        route.continue_()


async def block_resource_route_async(route, include_stylesheets: bool = False) -> None:
    if should_block_resource(route.request, include_stylesheets):
        await route.abort()
    else:
        await route.continue_()


# ── helpers ──────────────────────────────────────────────────────────────────

def safe_path_segment(value: str, fallback: str = "page") -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "-", value).strip(".-_")
    slug = re.sub(r"-+", "-", slug) or fallback
    return slug[:SAFE_PATH_SEGMENT_MAX].strip(".-_") or fallback


def path_text_length(path: Path) -> int:
    return len(str(path).replace("\\", "/"))


def url_to_filepath(url: str) -> Path:
    parsed = urlparse(url)
    url_hash = hashlib.sha1(url.encode("utf-8")).hexdigest()[:10]
    segments = [safe_path_segment(s) for s in parsed.path.strip("/").split("/") if s]
    if not segments:
        return Path(f"index-{url_hash}.pdf")

    filename_base = safe_path_segment(segments[-1], "page")
    filename = f"{filename_base}-{url_hash}.pdf"
    dirs = segments[:-1]
    rel_path = Path(*dirs, filename) if dirs else Path(filename)

    while dirs and path_text_length(rel_path) > SAFE_REL_PATH_MAX:
        dirs.pop(0)
        rel_path = Path(*dirs, filename) if dirs else Path(filename)

    if path_text_length(rel_path) > SAFE_REL_PATH_MAX:
        filename_base = filename_base[:max(16, SAFE_REL_PATH_MAX - len(url_hash) - 6)].strip(".-_") or "page"
        rel_path = Path(f"{filename_base}-{url_hash}.pdf")

    return rel_path


def load_manifest() -> dict:
    """Return manifest as a dict keyed by URL."""
    if MANIFEST_PATH.exists():
        with open(MANIFEST_PATH) as f:
            entries = json.load(f)
        return {e["url"]: e for e in entries}
    return {}


def save_manifest(manifest: dict) -> None:
    """Write manifest dict back to disk."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    with open(MANIFEST_PATH, "w") as f:
        json.dump(list(manifest.values()), f, indent=2)


def get_crawl_cache_path(start_url: str) -> Path:
    """Get the cache file path for a given starting URL (keyed by URL hash)."""
    import hashlib
    url_hash = hashlib.md5(start_url.encode()).hexdigest()[:8]
    CRAWL_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return CRAWL_CACHE_DIR / f"crawl_{url_hash}.json"


def load_crawl_cache(start_url: str) -> list | None:
    """Load cached crawl results for this URL, if they exist."""
    cache_path = get_crawl_cache_path(start_url)
    if cache_path.exists():
        try:
            with open(cache_path) as f:
                return json.load(f)
        except Exception:
            return None
    return None


def save_crawl_cache(start_url: str, links: list) -> None:
    """Save crawl results to cache."""
    cache_path = get_crawl_cache_path(start_url)
    with open(cache_path, "w") as f:
        json.dump(links, f, indent=2)


def _fetch_html(url: str) -> str:
    req = Request(url, headers={
        "User-Agent": "Mozilla/5.0 AppleWebKit/537.36 Chrome/120 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    })
    with urlopen(req, timeout=CRAWL_NAV_TIMEOUT / 1000) as response:
        content_type = response.headers.get("content-type", "")
        if "html" not in content_type.lower():
            return ""
        raw = response.read(2_000_000)
        charset = response.headers.get_content_charset() or "utf-8"
        return raw.decode(charset, errors="ignore")


def _extract_sublinks_from_html(markup: str, start_url: str, base: str, domain: str) -> set[str]:
    found = set()
    for match in re.finditer(r'''href\s*=\s*["']([^"'#]+)["']''', markup, flags=re.I):
        href = html.unescape(match.group(1)).strip()
        if not href or href.startswith(("mailto:", "tel:", "javascript:")):
            continue
        abs_url = urljoin(start_url, href)
        p2 = urlparse(abs_url)
        clean = f"{p2.scheme}://{p2.netloc}{p2.path}"
        clean_path = p2.path.rstrip("/")
        if "~/link/" in clean_path:
            continue
        if p2.netloc == domain and (clean_path == base or clean_path.startswith(base + "/")):
            found.add(clean)
    return found


def _scrape_worker(job_id: str, start_url: str, q: queue.Queue,
                   max_pages: int = CRAWL_MAX_PAGES, refresh: bool = False) -> None:
    """Background thread: crawl all sublinks under start_url (with caching)."""
    try:
        # Check if we have cached crawl results for this URL (unless a fresh
        # re-crawl was explicitly requested via refresh).
        cached_links = None if refresh else load_crawl_cache(start_url)
        if cached_links:
            q.put({"type": "using_cache", "count": len(cached_links)})
            manifest = load_manifest()
            links_with_status = [
                {
                    "url": l,
                    "done": l in manifest,
                    "file": manifest[l]["file"] if l in manifest else None,
                    "scraped_at": manifest[l]["scraped_at"] if l in manifest else None,
                }
                for l in cached_links
            ]
            emit_terminal(q, job_id, {"type": "complete", "links": links_with_status})
            return

        # No cache — crawl normally
        p0 = urlparse(start_url)
        base = p0.path.rstrip("/")
        domain = p0.netloc

        visited: set[str] = set()
        to_visit: list[str] = [start_url.rstrip("/") + "/"]
        all_found: set[str] = set()
        skipped = 0
        started_at = time.time()
        log_job(job_id, "crawl started", start_url=start_url, max_pages=max_pages, timeout_ms=CRAWL_NAV_TIMEOUT)
        update_job_status(job_id, stage="starting", current_url=start_url, visited=0, found=0, queued=1, skipped=0)

        while to_visit and len(visited) < max_pages:
            if should_stop(job_id):
                log_job(job_id, "crawl cancelled", visited=len(visited), found=len(all_found), skipped=skipped)
                emit_terminal(q, job_id, {"type": "cancelled", "reason": "Crawl cancelled."})
                return
            url = to_visit.pop(0)
            if url in visited:
                continue
            visited.add(url)

            elapsed = round(time.time() - started_at, 1)
            status = {"type": "crawling", "url": url, "count": len(visited),
                      "found": len(all_found), "queued": len(to_visit),
                      "skipped": skipped, "elapsed": elapsed, "stage": "loading"}
            update_job_status(job_id, stage="loading", current_url=url, visited=len(visited),
                              found=len(all_found), queued=len(to_visit), skipped=skipped, elapsed=elapsed)
            log_job(job_id, "loading page", visited=len(visited), found=len(all_found), queued=len(to_visit), url=url)
            q.put(status)

            try:
                markup = _fetch_html(url)
            except Exception as e:
                skipped += 1
                reason = str(e).splitlines()[0][:180]
                update_job_status(job_id, stage="skipped", current_url=url, visited=len(visited),
                                  found=len(all_found), queued=len(to_visit), skipped=skipped,
                                  last_error=reason, elapsed=round(time.time() - started_at, 1))
                log_job(job_id, "skipped page", url=url, reason=reason)
                q.put({"type": "skipped", "url": url, "count": len(visited),
                       "found": len(all_found), "queued": len(to_visit),
                       "skipped": skipped, "reason": reason})
                continue

            update_job_status(job_id, stage="extracting", current_url=url, visited=len(visited),
                              found=len(all_found), queued=len(to_visit), skipped=skipped,
                              elapsed=round(time.time() - started_at, 1))
            new_links = _extract_sublinks_from_html(markup, url, base, domain)
            all_found.update(new_links)
            log_job(job_id, "extracted links", url=url, new_links=len(new_links), total_found=len(all_found))

            for link in new_links:
                if link not in visited and link not in to_visit:
                    to_visit.append(link)
            update_job_status(job_id, stage="queued", current_url=url, visited=len(visited),
                              found=len(all_found), queued=len(to_visit), skipped=skipped,
                              elapsed=round(time.time() - started_at, 1))

        # Save crawl results to cache
        all_found_sorted = sorted(all_found)
        save_crawl_cache(start_url, all_found_sorted)

        manifest = load_manifest()
        links_with_status = [
            {
                "url": l,
                "done": l in manifest,
                "file": manifest[l]["file"] if l in manifest else None,
                "scraped_at": manifest[l]["scraped_at"] if l in manifest else None,
            }
            for l in all_found_sorted
        ]
        elapsed = round(time.time() - started_at, 1)
        log_job(job_id, "crawl complete", visited=len(visited), found=len(all_found_sorted), skipped=skipped, elapsed=elapsed)
        emit_terminal(q, job_id, {"type": "complete", "links": links_with_status,
                                  "visited": len(visited), "skipped": skipped, "elapsed": elapsed})
    except Exception as e:
        log_job(job_id, "crawl fatal", reason=str(e).splitlines()[0][:180])
        emit_terminal(q, job_id, {"type": "fatal", "reason": str(e)})



def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0563C1")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _build_docx_from_content(title, structured, out_path):
    doc = docx.Document()
    section = doc.sections[0]
    section.top_margin = Inches(1); section.bottom_margin = Inches(1)
    section.left_margin = Inches(1); section.right_margin = Inches(1)
    heading = doc.add_heading(title, level=1)
    if heading.runs:
        heading.runs[0].font.size = Pt(16)
    for item in (structured or []):
        if not isinstance(item, dict):
            continue
        tag = (item.get("tag") or "").lower()
        segments = []
        for raw in (item.get("runs") or []):
            if not isinstance(raw, dict):
                continue
            text = " ".join((raw.get("text") or "").split())
            if not text:
                continue
            href = (raw.get("href") or "").strip()
            if href.lower().startswith("javascript:"):
                href = ""
            segments.append((text, href))
        if not segments:
            continue
        if tag in ("h1", "h2"):
            para = doc.add_heading("", level=2); size = Pt(13)
        elif tag == "h3":
            para = doc.add_heading("", level=3); size = Pt(12)
        elif tag in ("h4", "h5", "h6"):
            para = doc.add_heading("", level=4); size = Pt(11)
        elif tag == "li":
            para = doc.add_paragraph(style="List Bullet"); size = Pt(11)
        else:
            para = doc.add_paragraph(); size = Pt(11)
        for idx, (text, href) in enumerate(segments):
            piece = text if idx == 0 else " " + text
            if href:
                try:
                    _add_hyperlink(para, piece, href)
                    continue
                except Exception:
                    pass
            run = para.add_run(piece)
            run.font.size = size
    doc.save(str(out_path))


def _download_worker(job_id: str, links: list[str], q: queue.Queue) -> None:

    manifest = load_manifest()
    pdf_paths = []
    docx_paths = []
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(args=["--no-sandbox", "--disable-setuid-sandbox"])
            context = None; page = None; render_count = 0
            def open_page():
                ctx = browser.new_context(user_agent="Mozilla/5.0 AppleWebKit/537.36 Chrome/120 Safari/537.36")
                ctx.route("**/*", lambda route: block_resource_route(route))
                return ctx, ctx.new_page()
            def close_page():
                if page: page.close()
                if context: context.close()
            context, page = open_page()
            try:
                for i, url in enumerate(links, 1):
                    if should_stop(job_id):
                        emit_terminal(q, job_id, {"type": "cancelled", "reason": "Download cancelled."})
                        return
                    if url in manifest:
                        q.put({"type": "skipped", "current": i, "total": len(links), "url": url, "file": manifest[url]["file"]})
                        continue
                    if render_count and render_count % PDF_PAGE_RECYCLE == 0:
                        close_page(); context, page = open_page()
                        q.put({"type": "status", "message": f"Recycled browser after {render_count} pages"})
                    update_job_status(job_id, stage="rendering", current=i, total=len(links), current_url=url, rendered=render_count)
                    q.put({"type": "progress", "current": i, "total": len(links), "url": url})
                    rel_pdf = url_to_filepath(url); out_pdf = OUTPUT_DIR / "pdf" / rel_pdf
                    rel_docx = url_to_filepath(url).with_suffix(".docx"); out_docx = OUTPUT_DIR / "docx" / rel_docx
                    out_pdf.parent.mkdir(parents=True, exist_ok=True)
                    out_docx.parent.mkdir(parents=True, exist_ok=True)
                    try:
                        page.emulate_media(media="screen")
                        page.goto(url, wait_until="domcontentloaded", timeout=PDF_NAV_TIMEOUT)
                        page.wait_for_timeout(PDF_SETTLE_MS)
                        page.add_style_tag(content=PAGE_CLEANUP_CSS)
                        page.pdf(path=str(out_pdf), format="A4", print_background=True, scale=0.9,
                                 margin={"top": "1cm", "bottom": "1cm", "left": "1cm", "right": "1cm"})
                        pdf_paths.append(out_pdf)
                        title = page.title().strip() or "Untitled"
                        structured = page.evaluate("""(boilerplate) => {
                            var results = [];
                            var skipTags = ['SCRIPT','STYLE','NOSCRIPT','HEADER','FOOTER','NAV','IMG','SVG','VIDEO','AUDIO','CANVAS','IFRAME','PICTURE','FIGURE'];
                            var blockTags = ['P','H1','H2','H3','H4','H5','H6','LI','DIV','SECTION','ARTICLE','MAIN','BLOCKQUOTE','TD','TH'];
                            function isSkipped(el) {
                                if (boilerplate && el.closest(boilerplate)) return true;
                                while (el) {
                                    if (skipTags.indexOf(el.tagName) !== -1) return true;
                                    el = el.parentElement;
                                }
                                return false;
                            }
                            function getBlock(el) {
                                while (el) {
                                    if (blockTags.indexOf(el.tagName) !== -1) return el;
                                    el = el.parentElement;
                                }
                                return null;
                            }
                            function getLink(el) {
                                while (el) {
                                    if (el.tagName === 'A' && el.href) {
                                        var h = el.getAttribute('href') || '';
                                        if (h.toLowerCase().indexOf('javascript:') === 0) return '';
                                        return el.href;
                                    }
                                    el = el.parentElement;
                                }
                                return '';
                            }
                            var walker = document.createTreeWalker(document.body, NodeFilter.SHOW_TEXT, null, false);
                            var node, currentBlock = null, pending = [];
                            function flush() {
                                var merged = [];
                                for (var i = 0; i < pending.length; i++) {
                                    var seg = pending[i];
                                    if (!seg.text) continue;
                                    var last = merged.length ? merged[merged.length - 1] : null;
                                    if (last && last.href === seg.href) {
                                        last.text = last.text + ' ' + seg.text;
                                    } else {
                                        merged.push({text: seg.text, href: seg.href});
                                    }
                                }
                                if (merged.length && currentBlock) {
                                    results.push({tag: currentBlock.tagName.toLowerCase(), runs: merged});
                                }
                                pending = [];
                            }
                            while (node = walker.nextNode()) {
                                var text = node.textContent.trim();
                                if (!text) continue;
                                var el = node.parentElement;
                                if (!el || isSkipped(el)) continue;
                                var cs = window.getComputedStyle(el);
                                if (cs && (cs.display === 'none' || cs.visibility === 'hidden')) continue;
                                var block = getBlock(el);
                                if (!block) continue;
                                if (block !== currentBlock) {
                                    flush();
                                    currentBlock = block;
                                }
                                pending.push({text: text, href: getLink(el)});
                            }
                            flush();
                            return results;
                        }""", BOILERPLATE_SELECTORS)
                        _build_docx_from_content(title, structured, out_docx)
                        docx_paths.append(out_docx)
                        render_count += 1
                        size_kb = round(out_pdf.stat().st_size / 1024, 1)
                        manifest[url] = {"url": url, "file": str(rel_pdf), "scraped_at": datetime.now().isoformat(timespec="seconds"), "status": "ok", "size_kb": size_kb}
                        save_manifest(manifest)
                        update_job_status(job_id, stage="done", current=i, total=len(links), current_url=url, rendered=render_count)
                        q.put({"type": "done_one", "url": url, "file": str(rel_pdf), "size_kb": size_kb})
                    except Exception as e:
                        q.put({"type": "error_one", "url": url, "reason": str(e)})
                    if i < len(links):
                        time.sleep(0.4)
            finally:
                close_page(); browser.close()
        if should_stop(job_id):
            emit_terminal(q, job_id, {"type": "cancelled", "reason": "Download cancelled."})
            return
        update_job_status(job_id, stage="zipping", rendered=render_count, total=len(links))
        q.put({"type": "status", "message": "Creating combined ZIP file..."})
        zip_path = OUTPUT_DIR / f"archive-{job_id}-{timestamp_slug()}.zip"
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for i2, p in enumerate(pdf_paths, 1):
                zf.write(p, p.relative_to(OUTPUT_DIR))
                if i2 % 30 == 0:
                    q.put({"type": "status", "message": f"Compressing ZIP... {i2}/{len(pdf_paths)}"})
            for i2, d in enumerate(docx_paths, 1):
                zf.write(d, d.relative_to(OUTPUT_DIR))
                if i2 % 30 == 0:
                    q.put({"type": "status", "message": f"Compressing ZIP... {len(pdf_paths) + i2}/{len(pdf_paths) + len(docx_paths)}"})
        jobs[job_id]["zip_path"] = str(zip_path)
        emit_terminal(q, job_id, {"type": "complete"})
    except Exception as e:
        emit_terminal(q, job_id, {"type": "fatal", "reason": str(e)})


def export_excel(links: list[dict], start_url: str) -> Path:
    """
    Build an Excel workbook from the scraped link list.
    Columns: #, Page Title (last path segment), Full URL, Folder Path, PDF File, Downloaded, Date, Size KB
    """
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Scraped Pages"

    # ── Styles ────────────────────────────────────────────────────────────────
    header_fill   = PatternFill("solid", fgColor="4F46E5")
    header_font   = Font(color="FFFFFF", bold=True, size=11)
    done_fill     = PatternFill("solid", fgColor="DCFCE7")
    pending_fill  = PatternFill("solid", fgColor="FEF9C3")
    center        = Alignment(horizontal="center", vertical="center")
    wrap          = Alignment(wrap_text=True, vertical="top")
    thin          = Side(style="thin", color="E5E7EB")
    border        = Border(left=thin, right=thin, top=thin, bottom=thin)

    # ── Header row ────────────────────────────────────────────────────────────
    headers = ["#", "Page Name", "Full URL", "Folder Path", "PDF File", "Downloaded", "Date Scraped", "Size (KB)"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    ws.row_dimensions[1].height = 22

    # ── Data rows ─────────────────────────────────────────────────────────────
    for i, link in enumerate(links, 1):
        url      = link.get("url", "")
        done     = link.get("done", False)
        file     = link.get("file") or ""
        scraped  = link.get("scraped_at") or ""
        size_kb  = link.get("size_kb") or ""

        page_name   = url_page_name(url)
        folder_path = url_parent_section(url)

        row_fill = done_fill if done else pending_fill
        row = [i, page_name, url, folder_path, file, "✓ Yes" if done else "✗ No", scraped, size_kb]

        for col, val in enumerate(row, 1):
            cell = ws.cell(row=i + 1, column=col, value=val)
            cell.fill = row_fill
            cell.border = border
            cell.alignment = wrap if col in (3, 4, 5) else center

    # ── Column widths ─────────────────────────────────────────────────────────
    widths = [5, 28, 60, 45, 50, 14, 20, 12]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(col)].width = w

    ws.freeze_panes = "A2"  # Keep header visible when scrolling

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet("Summary")
    total       = len(links)
    downloaded  = sum(1 for l in links if l.get("done"))
    pending     = total - downloaded

    ws2["A1"], ws2["B1"] = "Starting URL", start_url
    ws2["A2"], ws2["B2"] = "Total Pages Found", total
    ws2["A3"], ws2["B3"] = "Downloaded", downloaded
    ws2["A4"], ws2["B4"] = "Pending", pending
    ws2["A5"], ws2["B5"] = "Exported At", datetime.now().isoformat(timespec="seconds")

    for row in ws2.iter_rows(min_row=1, max_row=5, min_col=1, max_col=2):
        for cell in row:
            cell.border = border
            if cell.column == 1:
                cell.font = Font(bold=True)

    ws2.column_dimensions["A"].width = 22
    ws2.column_dimensions["B"].width = 70

    out_path = OUTPUT_DIR / f"scraped_pages_{uuid.uuid4().hex}.xlsx"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)
    return out_path


# ── Suggestion result cache (keyed by URL + guide version) ─────────────────────

def _suggest_cache_path(url: str) -> Path:
    """Cache file for an AI suggestion, invalidated when either source doc or the schema changes."""
    key = hashlib.md5(f"{url}|{SUGGEST_CACHE_VERSION}".encode()).hexdigest()[:12]
    SUGGEST_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return SUGGEST_CACHE_DIR / f"{key}.json"


def _load_suggest_cache(url: str) -> dict | None:
    path = _suggest_cache_path(url)
    if path.exists():
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except Exception:
            return None
    return None


def _save_suggest_cache(url: str, data: dict) -> None:
    try:
        _suggest_cache_path(url).write_text(json.dumps(data), encoding="utf-8")
    except Exception:
        pass


def url_page_name(url: str) -> str:
    """Human-readable page name derived from the last URL path segment."""
    segments = [s for s in urlparse(url).path.strip("/").split("/") if s]
    return segments[-1].replace("-", " ").title() if segments else "Home"


def url_parent_section(url: str) -> str:
    """Parent folder path derived from all but the last URL path segment."""
    segments = [s for s in urlparse(url).path.strip("/").split("/") if s]
    return "/".join(segments[:-1]) if len(segments) > 1 else "/"


# Worst rating across the six health dimensions wins the overall rating.
_HEALTH_SEVERITY = {
    "Strong": 1,
    "Acceptable": 2,
    "Needs Improvement": 3,
    "Significant Concern": 4,
}
_HEALTH_BY_SEVERITY = {v: k for k, v in _HEALTH_SEVERITY.items()}

_TREATMENT_FIT = {
    "KEEP": "Appropriate",
    "KEEP + IMPROVE": "Appropriate",
    "CONNECT": "Appropriate",
    "EXPERT REVIEW NEEDED": "Needs Expert Review",
}


def derive_computed_fields(url: str, fields: dict) -> dict:
    """
    Fields computed from the URL and the AI's own answers — no AI call needed.

    "Current Platform" is always the public Website because that is where every
    crawled page already lives.
    """
    treatment = fields.get("recommended_treatment", "")
    fit = _TREATMENT_FIT.get(treatment, "Reconsider" if treatment else "Unable to Determine")
    scored = [
        _HEALTH_SEVERITY[fields[dim]]
        for dim in HEALTH_DIMENSIONS
        if fields.get(dim) in _HEALTH_SEVERITY
    ]
    overall = _HEALTH_BY_SEVERITY[max(scored)] if scored else "Unable to Determine"
    return {
        "parent_section": url_parent_section(url),
        "current_platform": "Website",
        "current_platform_fit": fit,
        "overall_content_health": overall,
        "outdated_content_flag": fields.get("health_accuracy_currency") in ("Needs Improvement", "Significant Concern"),
        "conflicting_content_flag": fields.get("duplication_status") == "Conflicting content detected",
    }


def is_credit_error(exc: Exception) -> bool:
    text = str(exc).lower()
    return any(marker in text for marker in AI_CREDIT_ERROR_MARKERS)


def parse_ai_json(raw: str) -> dict:
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw.strip())
    match = re.search(r"\{.*\}", raw, flags=re.S)
    if not match:
        raise ValueError("No JSON object found")
    return json.loads(match.group(0))


def fallback_platform_suggestion(page_name: str, url: str, excerpt: str) -> dict:
    """
    Keyword-only guess used when the AI call fails after retries.

    Built on default_analysis_fields() so a fallback result carries the same keys
    as a real one — downstream code never has to special-case it.
    """
    text = f"{page_name} {url} {excerpt}".lower()
    if any(word in text for word in ("policy", "policies", "procedure", "form", "report", "complaint", "tobacco", "smoking", "violation")):
        guess = {
            "platform": "Maverick OneStop",
            "confidence": "Medium",
            "reason": "Fallback rule: policy, reporting, form, or procedure content usually belongs in Maverick OneStop.",
        }
    elif any(word in text for word in ("student organization", "club", "homecoming", "activities", "mavlife")):
        guess = {
            "platform": "MavLife / Student Hub",
            "confidence": "Low",
            "reason": "Fallback rule: student activity content may belong in MavLife / Student Hub.",
        }
    else:
        guess = {
            "platform": "Website",
            "confidence": "Low",
            "reason": "Fallback rule used because AI returned no valid classification; public-facing informational content defaults to Website.",
        }
    fields = default_analysis_fields()
    fields.update(guess)
    fields["confidence_gap"] = "Automated keyword fallback — a full AI analysis of this page is still needed."
    fields.update(derive_computed_fields(url, fields))
    return fields


async def _call_claude(client, page_name: str, url: str, excerpt: str) -> dict | None:
    """Classify one page with Claude. Retries transient API failures with backoff."""
    user_msg = (
        f"Page title: {page_name}\n"
        f"URL: {url}\n\n"
        f"Content excerpt:\n{excerpt}\n\n"
        "Analyze this page per the guide and the framework, and return the full JSON object. "
        "If the content is ambiguous or too thin to tell, use Low confidence."
    )
    last_error = "AI classification failed"
    for attempt in range(3):
        try:
            resp = await client.messages.create(
                model=TIER2_MODEL,
                max_tokens=TIER2_MAX_TOKENS,
                system=TIER2_SYSTEM_BLOCKS,
                messages=[{"role": "user", "content": user_msg}],
            )
        except Exception as e:
            if is_credit_error(e):
                raise AIQuotaError("Anthropic API credits appear to be exhausted or billing is unavailable. Add credits, then restart the analysis.") from e
            last_error = str(e)
            await asyncio.sleep(1.5 * (attempt + 1))
            continue
        try:
            raw = resp.content[0].text.strip()
            fields = clamp_analysis_fields(parse_ai_json(raw))
            fields.update(derive_computed_fields(url, fields))
            return fields
        except Exception as e:
            last_error = str(e)
            await asyncio.sleep(0.5 * (attempt + 1))
    fallback = fallback_platform_suggestion(page_name, url, excerpt)
    fallback["reason"] = f"{fallback['reason']} AI fallback reason: {last_error[:120]}."
    return fallback


async def _run_ai_analysis(job_id: str, results: list[dict], by_index: dict[int, dict], q: queue.Queue) -> None:
    """
    Fetch + classify every page, streaming upgrades as each finishes.

    Memory safety for large (~2000 link) runs:
      * Cached URLs never launch Chromium — they're emitted straight from disk.
      * Uncached URLs are processed in bounded batches; the browser is fully
        relaunched between batches so Chromium memory can't grow unbounded.
      * Concurrency within a batch is capped by ANALYSIS_CONCURRENCY.
    """
    total = len(results)
    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    done = 0

    def _emit(res: dict) -> None:
        """
        Stream one finished result (success or error) to the SSE queue.

        The whole result dict is spread into both the stored row and the SSE
        payload so new analysis fields flow through without being enumerated here.
        """
        idx = res["index"]
        stored = by_index[idx]
        if res.get("error"):
            stored.update(default_analysis_fields())
            stored.update(reason=res["error"], tier=0)
            q.put({"type": "analysis_error", "index": idx, "reason": res["error"],
                   "done": done, "total": total})
            return
        payload = {k: v for k, v in res.items() if k != "error"}
        stored.update(**payload, tier=2)
        q.put({"type": "analysis_upgrade", **payload,
               "cached": res.get("cached", False),
               "done": done, "total": total})

    # ── 1. Serve cached results instantly (no browser involved) ───────────────
    uncached: list[dict] = []
    for r in results:
        if should_stop(job_id):
            return
        cached = _load_suggest_cache(r["url"])
        if cached:
            done += 1
            _emit({"index": r["index"], **cached, "cached": True})
        else:
            uncached.append(r)

    if not uncached:
        return

    # ── 2. Process uncached URLs in browser-recycling batches ─────────────────
    sem = asyncio.Semaphore(ANALYSIS_CONCURRENCY)
    batch_size = max(ANALYSIS_BROWSER_RECYCLE, ANALYSIS_CONCURRENCY)

    async with async_playwright() as pw:
        for start in range(0, len(uncached), batch_size):
            if should_stop(job_id):
                return
            batch = uncached[start:start + batch_size]

            browser = await pw.chromium.launch(
                args=["--no-sandbox", "--disable-setuid-sandbox"])
            context = await browser.new_context(
                user_agent="Mozilla/5.0 AppleWebKit/537.36 Chrome/120 Safari/537.36"
            )
            await context.route("**/*", lambda route: block_resource_route_async(route, True))

            async def analyze_one(result: dict) -> dict:
                idx, url = result["index"], result["url"]
                if should_stop(job_id):
                    return {"index": idx, "error": "cancelled"}
                async with sem:
                    if should_stop(job_id):
                        return {"index": idx, "error": "cancelled"}
                    page = await context.new_page()
                    try:
                        await page.goto(url, wait_until="domcontentloaded",
                                        timeout=ANALYSIS_NAV_TIMEOUT)
                        title = (await page.title()).strip()
                        raw_text = await page.inner_text("body")
                    except Exception as e:
                        return {"index": idx, "error": str(e)}
                    finally:
                        await page.close()

                    excerpt = " ".join(raw_text.split()[:600])
                    page_name = title or result["page_name"]
                    ai = await _call_claude(client, page_name, url, excerpt)

                if ai is None:
                    ai = fallback_platform_suggestion(page_name, url, excerpt)
                out = {"page_name": page_name, **ai}
                _save_suggest_cache(url, out)
                return {"index": idx, **out}

            try:
                tasks = [asyncio.create_task(analyze_one(r)) for r in batch]
                for fut in asyncio.as_completed(tasks):
                    if should_stop(job_id):
                        for task in tasks:
                            task.cancel()
                        await asyncio.gather(*tasks, return_exceptions=True)
                        return
                    try:
                        res = await fut
                    except AIQuotaError:
                        for task in tasks:
                            task.cancel()
                        await asyncio.gather(*tasks, return_exceptions=True)
                        raise
                    if res.get("error") == "cancelled":
                        return
                    done += 1
                    _emit(res)
            finally:
                await context.close()
                await browser.close()


def _analysis_worker(job_id: str, links: list[dict], q: queue.Queue) -> None:
    """
    Background thread: AI-only platform classification.
      1. Emit instant placeholder rows so the table fills immediately.
      2. Concurrently fetch each page and let Claude classify it using the guide.
    """
    total = len(links)
    results: list[dict] = []
    by_index: dict[int, dict] = {}

    try:
        # ── Instant placeholder rows ──────────────────────────────────────────
        for i, link in enumerate(links, 1):
            if should_stop(job_id):
                emit_terminal(q, job_id, {"type": "cancelled", "reason": "Analysis cancelled."})
                return
            url = link.get("url", "")
            result = {
                **default_analysis_fields(),
                "index":      i,
                "total":      total,
                "url":        url,
                "page_name":  url_page_name(url),
                "platform":   "Analyzing…",
                "confidence": "—",
                "reason":     "Reading page content…",
                "tier":       1,
            }
            results.append(result)
            by_index[i] = result
            q.put({"type": "analysis_placeholder", **result})

        if not ANTHROPIC_API_KEY:
            emit_terminal(q, job_id, {"type": "fatal",
                                      "reason": "No ANTHROPIC_API_KEY set — add it to .env to enable AI analysis."})
            return

        # ── AI analysis (async, concurrent) ───────────────────────────────────
        q.put({"type": "analysis_started", "total": total})
        asyncio.run(_run_ai_analysis(job_id, results, by_index, q))

        if should_stop(job_id):
            emit_terminal(q, job_id, {"type": "cancelled", "reason": "Analysis cancelled."})
            return

        jobs[job_id]["analysis_results"] = results
        emit_terminal(q, job_id, {"type": "analysis_complete", "total": total})
    except AIQuotaError as e:
        emit_terminal(q, job_id, {"type": "fatal", "reason": str(e)})
    except Exception as e:
        emit_terminal(q, job_id, {"type": "fatal", "reason": str(e)})


def export_suggestions(results: list[dict], start_url: str) -> Path:
    """Build an Excel workbook from platform suggestion results."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Platform Suggestions"

    # ── Styles ────────────────────────────────────────────────────────────────
    header_fill = PatternFill("solid", fgColor="4F46E5")
    header_font = Font(color="FFFFFF", bold=True, size=11)
    center      = Alignment(horizontal="center", vertical="center")
    wrap        = Alignment(wrap_text=True, vertical="top")
    thin        = Side(style="thin", color="E5E7EB")
    border      = Border(left=thin, right=thin, top=thin, bottom=thin)

    PLATFORM_COLORS = {
        "Website":             "DBEAFE",   # blue-100
        "Maverick OneStop":    "FEF3C7",   # amber-100
        "The Fountain":        "EDE9FE",   # violet-100
        "MavLife / Student Hub": "D1FAE5", # green-100
        "Teams / SharePoint":  "CCFBF1",   # teal-100
        "No Clear Fit":        "F3F4F6",   # gray-100
    }
    CONFIDENCE_COLORS = {"High": "16A34A", "Medium": "D97706", "Low": "DC2626"}

    def _yes_no(value) -> str:
        return "Yes" if value else "No"

    def _joined(value) -> str:
        return "; ".join(value) if isinstance(value, list) else (value or "")

    # (header, value getter, column width, wrap text).  Grouped per the analysis
    # framework's suggested spreadsheet output.  Trailing owner/status columns
    # ship blank for the content team to fill in.
    COLUMNS = [
        # Content identification
        ("#",                             lambda r: r["index"],                              5,  False),
        ("Page Name",                     lambda r: r.get("page_name", ""),                  28, False),
        ("Full URL",                      lambda r: r.get("url", ""),                        50, True),
        ("Parent Section",                lambda r: r.get("parent_section", ""),             30, True),
        ("Current Platform",              lambda r: r.get("current_platform", "Website"),    18, False),
        ("Page Summary",                  lambda r: r.get("page_summary", ""),               50, True),
        # Audience and purpose
        ("Primary Audience",              lambda r: r.get("primary_audience", ""),           20, False),
        ("Secondary Audience",            lambda r: r.get("secondary_audience", ""),         20, False),
        ("Primary User Need",             lambda r: r.get("primary_user_need", ""),          16, False),
        ("Secondary User Need",           lambda r: r.get("secondary_user_need", ""),        16, False),
        ("Content Purpose",               lambda r: r.get("content_purpose", ""),            22, False),
        ("Purpose Explanation",           lambda r: r.get("purpose_explanation", ""),        40, True),
        # Content health
        ("Accuracy/Currency",             lambda r: r.get("health_accuracy_currency", ""),   18, False),
        ("Clarity",                       lambda r: r.get("health_clarity", ""),             18, False),
        ("Structure/Scanability",         lambda r: r.get("health_scanability_structure", ""), 18, False),
        ("Actionability",                 lambda r: r.get("health_actionability", ""),       18, False),
        ("User Focus",                    lambda r: r.get("health_user_focus", ""),          18, False),
        ("Sustainability",                lambda r: r.get("health_sustainability", ""),      18, False),
        ("Overall Content Health",        lambda r: r.get("overall_content_health", ""),     20, False),
        ("Content Health Notes",          lambda r: r.get("health_notes", ""),               50, True),
        # Content issues
        ("Outdated Content Flag",         lambda r: _yes_no(r.get("outdated_content_flag")), 16, False),
        ("Duplication Status",            lambda r: r.get("duplication_status", ""),         24, False),
        ("Related/Duplicate URLs",        lambda r: _joined(r.get("related_urls")),          45, True),
        ("Conflicting Content Flag",      lambda r: _yes_no(r.get("conflicting_content_flag")), 16, False),
        ("Major Issues Identified",       lambda r: r.get("major_issues", ""),               50, True),
        ("Improvement Opportunities",     lambda r: _joined(r.get("improvement_opportunities")), 45, True),
        # Strategy
        ("Current Platform Fit",          lambda r: r.get("current_platform_fit", ""),       18, False),
        ("Recommended Treatment",         lambda r: r.get("recommended_treatment", ""),      22, False),
        ("Recommended Platform",          lambda r: r.get("platform", ""),                   22, False),
        ("Platform Rationale",            lambda r: r.get("reason", ""),                     50, True),
        ("AI/Search Readiness",           lambda r: r.get("ai_search_readiness", ""),        18, False),
        # Project planning
        ("Priority",                      lambda r: r.get("priority", ""),                   24, False),
        ("Recommendation Confidence",     lambda r: r.get("confidence", ""),                 16, False),
        ("What Would Raise Confidence",   lambda r: r.get("confidence_gap", ""),             40, True),
        ("Key Recommendation",            lambda r: r.get("key_recommendation", ""),         50, True),
        ("Questions for Content Owner",   lambda r: _joined(r.get("content_owner_questions")), 45, True),
        ("Content Owner",                 lambda r: "",                                      20, False),
        ("Owner Decision",                lambda r: "",                                      20, False),
        ("Owner Notes",                   lambda r: "",                                      30, True),
        ("Status",                        lambda r: "",                                      16, False),
    ]

    for col, (header, _getter, _width, _wrap) in enumerate(COLUMNS, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    ws.row_dimensions[1].height = 22

    confidence_col = next(i for i, c in enumerate(COLUMNS, 1) if c[0] == "Recommendation Confidence")

    for r in results:
        row_num = r["index"] + 1
        row_fill = PatternFill("solid", fgColor=PLATFORM_COLORS.get(r.get("platform", ""), "F3F4F6"))
        conf_font = Font(color=CONFIDENCE_COLORS.get(r.get("confidence", ""), "374151"), bold=True)
        for col, (_header, getter, _width, wrap_text) in enumerate(COLUMNS, 1):
            cell = ws.cell(row=row_num, column=col, value=getter(r))
            cell.fill = row_fill
            cell.border = border
            cell.alignment = wrap if wrap_text else center
            if col == confidence_col:
                cell.font = conf_font

    for col, (_header, _getter, width, _wrap) in enumerate(COLUMNS, 1):
        ws.column_dimensions[get_column_letter(col)].width = width
    ws.freeze_panes = "D2"

    # ── Summary sheet ─────────────────────────────────────────────────────────
    ws2 = wb.create_sheet("Summary")
    from collections import Counter
    platform_counts = Counter(r["platform"] for r in results)
    conf_counts     = Counter(r["confidence"] for r in results)
    ws2["A1"], ws2["B1"] = "Starting URL", start_url
    ws2["A2"], ws2["B2"] = "Total Pages Analyzed", len(results)
    ws2["A3"], ws2["B3"] = "Exported At", datetime.now().isoformat(timespec="seconds")
    row = 5
    ws2.cell(row=row, column=1, value="Platform").font = Font(bold=True)
    ws2.cell(row=row, column=2, value="Count").font   = Font(bold=True)
    for platform, count in platform_counts.most_common():
        row += 1
        ws2.cell(row=row, column=1, value=platform)
        ws2.cell(row=row, column=2, value=count)
    row += 2
    ws2.cell(row=row, column=1, value="Confidence").font = Font(bold=True)
    ws2.cell(row=row, column=2, value="Count").font      = Font(bold=True)
    for conf, count in conf_counts.most_common():
        row += 1
        ws2.cell(row=row, column=1, value=conf)
        ws2.cell(row=row, column=2, value=count)

    for title, key in (("Recommended Treatment", "recommended_treatment"),
                       ("Priority", "priority"),
                       ("Overall Content Health", "overall_content_health")):
        row += 2
        ws2.cell(row=row, column=1, value=title).font = Font(bold=True)
        ws2.cell(row=row, column=2, value="Count").font = Font(bold=True)
        for value, count in Counter(r.get(key, "") for r in results).most_common():
            row += 1
            ws2.cell(row=row, column=1, value=value or "—")
            ws2.cell(row=row, column=2, value=count)

    ws2.column_dimensions["A"].width = 34
    ws2.column_dimensions["B"].width = 15

    out_path = OUTPUT_DIR / f"platform_suggestions_{uuid.uuid4().hex}.xlsx"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    wb.save(out_path)
    return out_path


# ── routes ───────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/healthz")
def healthz():
    return jsonify({"status": "ok"})


@app.route("/favicon.ico")
def favicon():
    return Response(status=204)


@app.route("/job-status/<job_id>")
def job_status(job_id: str):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify({
        "job_id": job_id,
        "type": job.get("type"),
        "created_at": job.get("created_at"),
        "finished_at": job.get("finished_at"),
        "cancelled": job.get("cancelled", False),
        "zip_ready": bool(job.get("zip_path")),
        "status": job.get("status", {}),
    })


@app.route("/cancel/<job_id>", methods=["POST"])
def cancel(job_id: str):
    if not cancel_job(job_id):
        return jsonify({"error": "Job not found or already finished"}), 404
    return jsonify({"status": "cancelling"})


@app.route("/scrape", methods=["POST"])
def scrape():
    body = request.json or {}
    url = body.get("url", "").strip()
    if not url:
        return jsonify({"error": "No URL provided"}), 400
    refresh = bool(body.get("refresh", False))

    job_id = str(uuid.uuid4())
    q: queue.Queue = queue.Queue()
    if not register_job(job_id, {"type": "scrape", "queue": q, "links": None, "zip_path": None}):
        return jsonify({"error": "Too many active jobs. Wait for the current job to finish, then try again."}), 429
    threading.Thread(target=_scrape_worker, args=(job_id, url, q),
                     kwargs={"refresh": refresh}, daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/start-download", methods=["POST"])
def start_download():
    links = (request.json or {}).get("links", [])
    if not links:
        return jsonify({"error": "No links provided"}), 400
    job_id = str(uuid.uuid4())
    q: queue.Queue = queue.Queue()
    if not register_job(job_id, {"type": "download", "queue": q, "zip_path": None}):
        return jsonify({"error": "Too many active jobs. Wait for the current job to finish, then try again."}), 429
    threading.Thread(target=_download_worker, args=(job_id, links, q), daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/progress/<job_id>")
def progress(job_id: str):
    job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404

    def stream():
        q = job["queue"]
        replay = job.get("terminal")
        if replay:
            if replay.get("type") == "complete" and job.get("type") == "scrape":
                job["links"] = replay.get("links", [])
            if replay.get("type") == "analysis_complete":
                analysis_results = job.get("analysis_results") or []
                for r in analysis_results:
                    placeholder = {"type": "analysis_placeholder", **r}
                    upgrade = {"type": "analysis_upgrade", **r,
                               "done": len(analysis_results), "total": len(analysis_results)}
                    yield f"data: {json.dumps(placeholder)}\n\n"
                    yield f"data: {json.dumps(upgrade)}\n\n"
            # The terminal event last, mirroring live ordering, so a reconnected
            # client still runs its completion handling instead of hanging.
            yield f"data: {json.dumps(replay)}\n\n"
            return
        while True:
            try:
                evt = q.get(timeout=SSE_HEARTBEAT_SECONDS)
                if evt.get("type") == "complete" and job.get("type") == "scrape":
                    job["links"] = evt.get("links", [])
                yield f"data: {json.dumps(evt)}\n\n"
                if evt.get("type") in ("complete", "fatal", "cancelled"):
                    job["finished_at"] = time.time()
                    break
            except queue.Empty:
                yield 'data: {"type":"heartbeat"}\n\n'

    return Response(stream(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.route("/get-zip/<job_id>")
def get_zip(job_id: str):
    job = jobs.get(job_id)
    if not job or not job.get("zip_path"):
        return "Not ready", 404
    return send_file(job["zip_path"], as_attachment=True, download_name=f"pages_{timestamp_slug()}.zip")


@app.route("/export-excel", methods=["POST"])

def export_excel_route():
    """Generate and return an Excel file from the current link list + manifest."""
    data       = request.json or {}
    links      = data.get("links", [])
    start_url  = data.get("start_url", "")
    if not links:
        return jsonify({"error": "No links provided"}), 400
    try:
        # Merge manifest status into links in case it's fresher than what the UI has
        manifest = load_manifest()
        for link in links:
            url = link.get("url", "")
            if url in manifest:
                link["done"]       = True
                link["file"]       = manifest[url]["file"]
                link["scraped_at"] = manifest[url]["scraped_at"]
                link["size_kb"]    = manifest[url].get("size_kb", "")
        path = export_excel(links, start_url)
        return send_file(str(path), as_attachment=True, download_name=f"scraped_pages_{timestamp_slug()}.xlsx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/analyze-platforms", methods=["POST"])
def analyze_platforms():
    """Start a background platform-suggestion job for the given link list."""
    data  = request.json or {}
    links = data.get("links", [])
    if not links:
        return jsonify({"error": "No links provided"}), 400
    job_id = str(uuid.uuid4())
    q: queue.Queue = queue.Queue()
    if not register_job(job_id, {"type": "analysis", "queue": q, "analysis_results": None}):
        return jsonify({"error": "Too many active jobs. Wait for the current job to finish, then try again."}), 429
    threading.Thread(target=_analysis_worker, args=(job_id, links, q), daemon=True).start()
    return jsonify({"job_id": job_id})


@app.route("/export-suggestions", methods=["POST"])
def export_suggestions_route():
    """Generate and return an Excel file from the platform suggestion results."""
    data      = request.json or {}
    results   = data.get("results", [])
    start_url = data.get("start_url", "")
    if not results:
        return jsonify({"error": "No results provided"}), 400
    try:
        path = export_suggestions(results, start_url)
        return send_file(str(path), as_attachment=True, download_name=f"platform_suggestions_{timestamp_slug()}.xlsx")
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/clear-manifest", methods=["POST"])
def clear_manifest():
    """Clear the manifest to start fresh."""
    try:
        if MANIFEST_PATH.exists():
            MANIFEST_PATH.unlink()
        return jsonify({"status": "cleared"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(
        debug=os.getenv("FLASK_DEBUG") == "1",
        host="0.0.0.0",
        port=env_int("PORT", 5000, 1),
        threaded=True,
    )
